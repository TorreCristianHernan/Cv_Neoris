from playwright.sync_api import sync_playwright
from docx import Document
from bs4 import BeautifulSoup
from tkinter import Tk, Label, Entry, Button, filedialog

def download_word(url):
    """
    Downloads the visible text content from a given URL and saves it as a Word document.

    Args:
        url (str): The URL to download the content from.

    Returns:
        None

    Raises:
        None
    """
    # Launch the browser
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()

        # Navigate to the URL
        page.goto(url)

        # Wait for the page to load completely
        page.wait_for_load_state("networkidle")
        
        print("after loading, before scroll")
        page.evaluate("document.querySelector('a[title=\"Use ALT+1, ALT+SHIFT+1, or CMD+SHIFT+1 to go to the main content.\"]').scrollIntoViewIfNeeded()")
        print("after scrolling")
        page.screenshot(path="screenshot.png")

        input_element = page.locator('#__input0-inner')
        print("after selecting the input")
        input_element.fill('neorisconsP2')
        print("after entering info input")
        page.screenshot(path="screenshot2.png")
        input_element.press('Enter')

        # button_element = page.locator("#__button0")
        # print("after selecting button")
        # button_element.click()
        # print("after click button")
        page.wait_for_timeout(15000)
        #
        page.screenshot(path="screenshot3.png")

        #      
        # Get the page content
        page_content = page.content()
        print("print", page_content)

        # Parse the HTML content using BeautifulSoup (optional)
        soup = BeautifulSoup(page_content, 'html.parser')
        
        # Create a new Word document
        document = Document()
        
        ##
        document.add_paragraph(page_content)
        ##
        # # Find all elements that contain visible text
        # for element in soup.find_all(string=True):
        #     if element.parent.name not in ['script', 'style']:
        #         document.add_paragraph(str(element))

        # Save the Word document
        file_path = filedialog.asksaveasfilename(defaultextension=".docx")
        document.save(file_path)
        print("The Word document has been created successfully.")

        # Close the browser
        browser.close()

# Create the Tkinter window
window = Tk()
window.title("Download Web Page as Word Document")

# Create a label and entry for the URL input
url_label = Label(window, text="URL:")
url_label.pack()
url_entry = Entry(window)
url_entry.pack()

# Create a button to initiate the download
download_button = Button(window, text="Download as Word", command=lambda: download_word(url_entry.get()))
download_button.pack()

# Run the Tkinter event loop
window.mainloop()