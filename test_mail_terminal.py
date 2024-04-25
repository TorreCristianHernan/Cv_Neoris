import unittest
from unittest.mock import patch
from io import StringIO
from bs4 import BeautifulSoup
from docx import Document

# Import the function to test
from mail_terminal import download_word

class TestDownloadWord(unittest.TestCase):
    @patch('requests.get')
    @patch('builtins.input', return_value='example.com')
    @patch('builtins.print')
    @patch('filedialog.asksaveasfilename', return_value='test.docx')
    def test_download_word_success(self, mock_asksaveasfilename, mock_print, mock_input, mock_get):
        # Mock the response from requests.get
        mock_response = mock_get.return_value
        mock_response.status_code = 200
        mock_response.text = '<html><body><p>Test content</p></body></html>'

        # Mock the BeautifulSoup object
        mock_soup = BeautifulSoup(mock_response.text, 'html.parser')

        # Mock the Document object
        mock_document = Document()

        # Call the function to test
        download_word('example.com')

        # Assertions
        mock_get.assert_called_once_with('example.com')
        mock_soup.find_all.assert_called_once_with(string=True)
        mock_document.add_paragraph.assert_called_once_with('Test content')
        mock_document.save.assert_called_once_with('test.docx')
        mock_print.assert_called_once_with('The Word document has been created successfully.')

    @patch('requests.get')
    @patch('builtins.input', return_value='example.com')
    @patch('builtins.print')
    def test_download_word_error(self, mock_print, mock_input, mock_get):
        # Mock the response from requests.get
        mock_response = mock_get.return_value
        mock_response.status_code = 404

        # Call the function to test
        download_word('example.com')

        # Assertions
        mock_get.assert_called_once_with('example.com')
        mock_print.assert_called_once_with('Error accessing the page: Status 404')

if __name__ == '__main__':
    unittest.main()import unittest
from unittest.mock import patch
from io import StringIO
from bs4 import BeautifulSoup
from docx import Document
import requests
from tkinter import filedialog

# Import the function to test
from mail_terminal import download_word

class TestDownloadWord(unittest.TestCase):
    @patch('requests.get')
    @patch('builtins.input', return_value='example.com')
    @patch('builtins.print')
    @patch('filedialog.asksaveasfilename', return_value='test.docx')
    def test_download_word_success(self, mock_asksaveasfilename, mock_print, mock_input, mock_get):
        # Mock the response from requests.get
        mock_response = mock_get.return_value
        mock_response.status_code = 200
        mock_response.text = '<html><body><p>Test content</p></body></html>'

        # Mock the BeautifulSoup object
        mock_soup = BeautifulSoup(mock_response.text, 'html.parser')

        # Mock the Document object
        mock_document = Document()

        # Call the function to test
        download_word('example.com')

        # Assertions
        mock_get.assert_called_once_with('example.com')
        mock_soup.find_all.assert_called_once_with(string=True)
        mock_document.add_paragraph.assert_called_once_with('Test content')
        mock_document.save.assert_called_once_with('test.docx')
        mock_print.assert_called_once_with('The Word document has been created successfully.')

    @patch('requests.get')
    @patch('builtins.input', return_value='example.com')
    @patch('builtins.print')
    def test_download_word_error(self, mock_print, mock_input, mock_get):
        # Mock the response from requests.get
        mock_response = mock_get.return_value
        mock_response.status_code = 404

        # Call the function to test
        download_word('example.com')

        # Assertions
        mock_get.assert_called_once_with('example.com')
        mock_print.assert_called_once_with('Error accessing the page: Status 404')

if __name__ == '__main__':
    unittest.main()